"""文件模板服务 —— 下载时自动替换 {{占位符}} 为实例实际值

支持格式：
  - .docx：python-docx 遍历段落/表格/页眉页脚替换
  - .xlsx：openpyxl 遍历所有单元格替换

关键处理：Word 可能将 {{变量}} 拆分到多个 Run 中，需要拼接后统一替换。
"""

import logging
import re
import os
from datetime import date
from io import BytesIO

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models import (
    DocumentTemplate, FlowTemplate, FlowInstance, Task,
    InstanceNode, User, Organization,
)

logger = logging.getLogger(__name__)

# 占位符 → 实例字段映射（中文变量名，管理员在模板中使用）
# 格式：{{变量名}}
_VAR_MAP: dict[str, tuple[str, str]] = {
    "项目名称": ("instance.name", "str"),
    "项目描述": ("instance.description", "str"),
    "合同号": ("instance.contract_no", "str"),
    "产品型号": ("instance.product_model", "str"),
    "销售经理": ("instance.sales_manager", "str"),
    "模板名称": ("instance.template_name", "str"),
    "优先级": ("instance.priority", "str"),
    "当前节点": ("node.name", "str"),
    "发起人": ("initiator.real_name", "str"),
    "发起日期": ("instance.initiated_at", "date"),
    "所属部门": ("org.name", "str"),
    "当前负责人": ("assignee.real_name", "str"),
    "当前日期": ("today", "date"),
}

# 收集所有合法的占位符 key
_ALL_PLACEHOLDERS = set(f"{{{{{k}}}}}" for k in _VAR_MAP.keys())


async def resolve_template_variables(
    db: AsyncSession,
    doc_template_id: int,
    task_id: int,
) -> dict[str, str]:
    """解析模板变量为实际值。

    根据文档模板找到所属流程模板，再通过 Task 找到流程实例和相关人员，
    逐一解析 _VAR_MAP 中的变量为实际字符串值。

    返回：{"{{变量名}}": "实际值", ...}
    """

    # 1. 查文档模板
    doc = (await db.execute(
        select(DocumentTemplate).where(DocumentTemplate.id == doc_template_id)
    )).scalar_one_or_none()
    if doc is None:
        raise ValueError(f"文件模板不存在: {doc_template_id}")

    # 2. 查 Task → 节点 → 实例（手动查询关联对象，模型没有定义 relationship）
    task = (await db.execute(
        select(Task).where(Task.id == task_id)
    )).scalar_one_or_none()
    if task is None:
        raise ValueError(f"任务不存在: {task_id}")

    # 手动查负责人 User
    assignee = (await db.execute(
        select(User).where(User.id == task.assignee_id)
    )).scalar_one_or_none()

    node = (await db.execute(
        select(InstanceNode).where(InstanceNode.id == task.node_id)
    )).scalar_one_or_none()

    instance = (await db.execute(
        select(FlowInstance).where(FlowInstance.id == task.instance_id)
    )).scalar_one_or_none()
    if instance is None:
        raise ValueError(f"流程实例不存在: {task.instance_id}")

    # 通过实例获取流程模板（文档模板经中间表关联，此处用实例的模板上下文）
    template = (await db.execute(
        select(FlowTemplate).where(FlowTemplate.id == instance.template_id)
    )).scalar_one_or_none()

    # 手动查发起人 User
    initiator = (await db.execute(
        select(User).where(User.id == instance.initiator_id)
    )).scalar_one_or_none()

    # 手动查发起人所在组织
    org = None
    if initiator and initiator.organization_id:
        org = (await db.execute(
            select(Organization).where(Organization.id == initiator.organization_id)
        )).scalar_one_or_none()

    # 3. 解析每个变量
    result: dict[str, str] = {}
    today_str = date.today().strftime("%Y-%m-%d")

    for var_name, (field_path, field_type) in _VAR_MAP.items():
        placeholder = f"{{{{{var_name}}}}}"

        if field_path == "today":
            result[placeholder] = today_str
            continue

        # 按路径解析
        parts = field_path.split(".")
        value = None

        # 从各数据源获取值
        if parts[0] == "instance":
            value = getattr(instance, parts[1], None)
        elif parts[0] == "template":
            value = getattr(template, parts[1], None) if template else None
        elif parts[0] == "node":
            value = getattr(node, parts[1], None) if node else None
        elif parts[0] == "initiator":
            value = getattr(initiator, parts[1], None) if initiator else None
        elif parts[0] == "org":
            value = getattr(org, parts[1], None) if org else None
        elif parts[0] == "assignee":
            value = getattr(assignee, parts[1], None) if assignee else None

        if value is None:
            result[placeholder] = ""
        elif field_type == "date":
            # 日期类型格式化为 YYYY-MM-DD
            if hasattr(value, "strftime"):
                result[placeholder] = value.strftime("%Y-%m-%d")
            else:
                result[placeholder] = str(value)[:10]
        else:
            result[placeholder] = str(value)

    return result


def _replace_in_runs(paragraph, replacements: dict[str, str]) -> bool:
    """替换段落中（可能跨多个 Run）的占位符文本。

    Word 经常将连续文本拆分到不同的 Run 中，所以需要：
      1. 拼接所有 Run 文本
      2. 在拼接文本中找到所有 {{变量}}
      3. 逐一替换为实际值
      4. 将结果重新写回 Run 列表

    返回是否有替换发生。
    """
    if not paragraph.runs:
        return False

    # 拼接全文
    full_text = "".join(r.text for r in paragraph.runs)
    if not full_text:
        return False

    # 查找所有占位符
    found = set(re.findall(r"\{\{[一-龥a-zA-Z0-9]+\}\}", full_text))
    if not found:
        return False

    replaced = False
    for placeholder in found:
        if placeholder in replacements:
            full_text = full_text.replace(placeholder, replacements[placeholder])
            replaced = True

    if not replaced:
        return False

    # 清除所有 Run，将结果写入第一个 Run
    for run in paragraph.runs[1:]:
        run.text = ""
    paragraph.runs[0].text = full_text
    return True


def _replace_in_table(table, replacements: dict[str, str]):
    """替换表格中所有单元格的占位符"""
    for row in table.rows:
        for cell in row.cells:
            # 单元格内可能有多个段落
            for paragraph in cell.paragraphs:
                _replace_in_runs(paragraph, replacements)


def _replace_in_docx(doc, replacements: dict[str, str]):
    """替换 Word 文档中所有位置的占位符：正文段落 + 表格 + 页眉页脚"""
    # 正文段落
    for paragraph in doc.paragraphs:
        _replace_in_runs(paragraph, replacements)

    # 表格
    for table in doc.tables:
        _replace_in_table(table, replacements)

    # 页眉页脚
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_runs(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_runs(paragraph, replacements)


def fill_docx_template(file_path: str, replacements: dict[str, str]) -> BytesIO:
    """加载 .docx 模板，替换占位符，返回内存中的文件流"""
    from docx import Document

    doc = Document(file_path)
    _replace_in_docx(doc, replacements)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def fill_xlsx_template(file_path: str, replacements: dict[str, str]) -> BytesIO:
    """加载 .xlsx 模板，替换所有单元格中的占位符，返回内存中的文件流"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    for placeholder, value in replacements.items():
                        if placeholder in cell.value:
                            cell.value = cell.value.replace(placeholder, value)

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def fill_template(file_path: str, file_type: str, replacements: dict[str, str]) -> BytesIO:
    """根据文件类型调用对应的填充函数"""
    if file_type == "xlsx":
        return fill_xlsx_template(file_path, replacements)
    else:
        # 默认走 docx
        return fill_docx_template(file_path, replacements)


def get_doc_template_abs_path(doc: DocumentTemplate) -> str:
    """获取文件模板的绝对路径（兼容存储各种路径格式）"""
    if os.path.isabs(doc.file_path):
        path = doc.file_path
        # 绝对路径：直接返回
        if os.path.exists(path):
            return path
    else:
        # 相对路径：先尝试与 STORAGE_ROOT 拼接
        path = os.path.join(settings.STORAGE_ROOT, doc.file_path)
        if os.path.exists(path):
            return path
        # 如果 file_path 已包含 STORAGE_ROOT 目录名（如 "storage/doc_templates/..."），
        # 避免重复拼接 —— 尝试直接作为相对路径返回
        if os.path.exists(doc.file_path):
            return doc.file_path
    # 最终兜底（确保返回绝对路径）
    return os.path.abspath(path)
