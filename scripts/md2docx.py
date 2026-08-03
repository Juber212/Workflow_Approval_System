"""将用户使用手册 Markdown 转换为专业 Word 文档（含封面 + 自动目录）"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
# 尝试导入 PIL 读取图片尺寸
try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_ROOT / "docs/user-manual/用户使用手册.md"
OUT_PATH = PROJECT_ROOT / "docs/user-manual/用户使用手册.docx"

FONT_BODY = "微软雅黑"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_H1 = Pt(22)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_H4 = Pt(11.5)
FONT_SIZE_TABLE = Pt(9.5)
COLOR_PRIMARY = RGBColor(0x1A, 0x6F, 0xB5)   # 主题蓝
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_CAPTION = RGBColor(0x99, 0x99, 0x99)
COLOR_TABLE_HEADER_BG = "D6E8F7"
COLOR_TABLE_BORDER = "B0C4DE"


def strip_markdown(text: str) -> str:
    """去掉行内 markdown 标记：**bold** `code` *italic*"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text


def set_run_font(run, name=FONT_BODY, size=FONT_SIZE_BODY, color=COLOR_BODY, bold=False, italic=False):
    """统一设置 run 字体"""
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_styled_paragraph(doc, text, style=None, font_size=None, color=None, bold=False, italic=False,
                         alignment=None, space_before=None, space_after=None, first_line_indent=None):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    if style:
        p.style = style

    # 解析内联格式: **bold** 和 `code`
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=font_size or FONT_SIZE_BODY, color=color or COLOR_BODY, bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_run_font(run, name=FONT_CODE, size=font_size or Pt(9.5), color=color or COLOR_BODY)
        else:
            run = p.add_run(part)
            set_run_font(run, size=font_size or FONT_SIZE_BODY, color=color or COLOR_BODY, bold=bold, italic=italic)

    if alignment is not None:
        p.alignment = alignment

    pf = p.paragraph_format
    pf.line_spacing = 1.5
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = first_line_indent

    return p


def setup_heading_styles(doc):
    """修改 Word 内置标题样式的外观（不改名称，TOC 才能识别）"""
    heading_styles = {
        'Heading 1': {'size': FONT_SIZE_H1, 'color': COLOR_HEADING, 'space_before': Pt(32), 'space_after': Pt(12), 'border_color': '1A6FB5', 'border_sz': '4'},
        'Heading 2': {'size': FONT_SIZE_H2, 'color': COLOR_HEADING, 'space_before': Pt(24), 'space_after': Pt(8), 'border_color': 'CCCCCC', 'border_sz': '2'},
        'Heading 3': {'size': FONT_SIZE_H3, 'color': COLOR_HEADING, 'space_before': Pt(18), 'space_after': Pt(6), 'border_color': None, 'border_sz': None},
        'Heading 4': {'size': FONT_SIZE_H4, 'color': COLOR_HEADING, 'space_before': Pt(14), 'space_after': Pt(4), 'border_color': None, 'border_sz': None},
    }
    for style_name, cfg in heading_styles.items():
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        font = style.font
        font.name = FONT_BODY
        font.size = cfg['size']
        font.color.rgb = cfg['color']
        font.bold = True
        style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        pf = style.paragraph_format
        pf.line_spacing = 1.5
        pf.space_before = cfg['space_before']
        pf.space_after = cfg['space_after']
        pf.keep_with_next = True
        # 底部边框
        if cfg['border_color']:
            pPr = style.element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): cfg['border_sz'],
                qn('w:space'): '4',
                qn('w:color'): cfg['border_color'],
            })
            pBdr.append(bottom)
            pPr.append(pBdr)


def add_heading_styled(doc, text, level):
    """使用 Word 内置标题样式（TOC 可识别）"""
    h = doc.add_heading(text, level=level)
    # 修正中文字体（add_heading 有时不继承样式字体）
    for run in h.runs:
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)


def set_cell_text(cell, text, bold=False, font_size=FONT_SIZE_TABLE):
    """设置表格单元格文本（已去除 markdown）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(strip_markdown(text))
    set_run_font(run, size=font_size, bold=bold, color=COLOR_BODY)
    return p


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows):
    """添加美化的表格"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.autofit = True

    # 设置表格边框
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{COLOR_TABLE_BORDER}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True)
        set_cell_shading(cell, COLOR_TABLE_HEADER_BG)

    # 数据行（交替底色）
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            set_cell_text(cell, val)
            if r % 2 == 1:
                set_cell_shading(cell, "F5F8FC")

    doc.add_paragraph()  # 表后空行


IMAGES_DIR = PROJECT_ROOT / "docs/user-manual/images"


def add_image_placeholder(doc, alt, src):
    """插入图片：文件存在则嵌入，否则显示占位符"""
    img_path = IMAGES_DIR / Path(src).name

    if img_path.exists():
        # 获取图片尺寸（PIL 可用时自动计算，否则用默认宽高）
        max_width = Inches(5.5)
        if HAS_PIL:
            try:
                with PILImage.open(str(img_path)) as im:
                    w_px, h_px = im.size
                    ratio = h_px / w_px if w_px > 0 else 0.75
                    width = max_width
                    height = width * ratio
            except Exception:
                width, height = max_width, Inches(3.1)
        else:
            width, height = max_width, Inches(3.1)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(str(img_path), width=width, height=height)

        # 图注
        add_styled_paragraph(
            doc, f"△ {alt}",
            font_size=Pt(8.5), color=COLOR_CAPTION, italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(0), space_after=Pt(12),
        )
    else:
        # 占位符
        p = add_styled_paragraph(
            doc, f"[ 截图：{alt} ]",
            font_size=Pt(9), color=COLOR_CAPTION, italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(8), space_after=Pt(2),
        )
        add_styled_paragraph(
            doc, f"（待补：{src}）",
            font_size=Pt(7.5), color=RGBColor(0xBB, 0xBB, 0xBB),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(0), space_after=Pt(8),
        )


def add_blockquote(doc, text):
    """引用块"""
    p = add_styled_paragraph(
        doc, strip_markdown(text),
        font_size=Pt(9.5), color=RGBColor(0x66, 0x66, 0x66), italic=True,
        space_before=Pt(4), space_after=Pt(4),
    )
    p.paragraph_format.left_indent = Cm(0.8)
    # 左边框
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    left = pBdr.makeelement(qn('w:left'), {
        qn('w:val'): 'single',
        qn('w:sz'): '12',
        qn('w:space'): '8',
        qn('w:color'): '1A6FB5',
    })
    pBdr.append(left)
    pPr.append(pBdr)


def add_hr(doc):
    """水平分割线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): 'CCCCCC',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def convert():
    doc = Document()

    # ── 全局默认样式 ──
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = FONT_SIZE_BODY
    style.font.color.rgb = COLOR_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)

    # 修改内置标题样式（必须在添加内容之前）
    setup_heading_styles(doc)

    # ── 封面 ──
    # 空行推到中间
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("企业流程审批系统")
    set_run_font(run, size=Pt(32), color=COLOR_PRIMARY, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("用户使用手册")
    set_run_font(run2, size=Pt(24), color=COLOR_HEADING, bold=True)

    # 分隔线
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(16)
    p3.paragraph_format.space_after = Pt(16)
    run3 = p3.add_run("━" * 30)
    set_run_font(run3, size=Pt(8), color=COLOR_PRIMARY)

    # 版本信息
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("版本 1.0  |  2026 年 7 月")
    set_run_font(run4, size=Pt(11), color=COLOR_CAPTION)

    # 分页
    doc.add_page_break()

    # ── 目录页 ──
    p_toc_title = doc.add_paragraph()
    p_toc_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_toc = p_toc_title.add_run("目  录")
    set_run_font(run_toc, size=FONT_SIZE_H1, color=COLOR_HEADING, bold=True)
    p_toc_title.paragraph_format.space_after = Pt(16)

    # 插入 Word 目录域（需在 Word 中右键更新域）
    paragraph = doc.add_paragraph()
    run_toc_field = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_toc_field._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">TOC \\o "1-3" \\h \\z \\u</w:instrText>')
    run_toc_field._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_toc_field._element.append(fldChar2)
    run_toc_text = paragraph.add_run("（打开 Word 后右键此处 → 更新域，自动生成目录）")
    set_run_font(run_toc_text, size=Pt(9), color=COLOR_CAPTION, italic=True)
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_toc_field._element.append(fldChar3)

    doc.add_page_break()

    # ── 正文 ──
    text = MD_PATH.read_text(encoding='utf-8')
    lines = text.split('\n')

    # 跳过第一行的大标题（封面已有）
    skip_main_title = True

    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过文件开头主标题和紧接的链接/空行
        if skip_main_title and (line.startswith('# ') or not line.strip() or line.startswith('>')):
            if line.startswith('# '):
                skip_main_title = False  # 只跳第一个 h1
            i += 1
            continue
        skip_main_title = False

        # 空行
        if not line.strip():
            i += 1
            continue

        # 图片 ![alt](src)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', line.strip())
        if img_match:
            add_image_placeholder(doc, img_match.group(1), img_match.group(2))
            i += 1
            continue

        # 水平线 ---
        if re.match(r'^-{3,}$', line.strip()):
            add_hr(doc)
            i += 1
            continue

        # 表格
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            content_rows = []
            for tl in table_lines:
                if re.match(r'^\|[\s\-:|]+\|$', tl.strip()):
                    continue
                cells = [c.strip() for c in tl.strip().split('|')[1:-1]]
                content_rows.append(cells)
            if len(content_rows) >= 2:
                add_table(doc, content_rows[0], content_rows[1:])
            continue

        # 标题 h1-h4
        heading_match = re.match(r'^(#{1,4})\s+(.+)', line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            add_heading_styled(doc, heading_text, level)
            i += 1
            continue

        # 引用块 >
        if line.strip().startswith('>'):
            add_blockquote(doc, line.strip()[1:].strip())
            i += 1
            continue

        # 无序列表
        ul_match = re.match(r'^(\s*)-\s+(.+)', line)
        if ul_match:
            add_styled_paragraph(doc, strip_markdown(ul_match.group(2)), space_before=Pt(1))
            # 给列表项加左缩进
            doc.paragraphs[-1].paragraph_format.left_indent = Cm(0.8)
            doc.paragraphs[-1].paragraph_format.first_line_indent = Cm(-0.4)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)', line)
        if ol_match:
            add_styled_paragraph(doc, strip_markdown(ol_match.group(2)), space_before=Pt(1))
            doc.paragraphs[-1].paragraph_format.left_indent = Cm(0.8)
            doc.paragraphs[-1].paragraph_format.first_line_indent = Cm(-0.4)
            i += 1
            continue

        # 普通段落
        add_styled_paragraph(doc, line.strip())
        i += 1

    # 保存
    doc.save(str(OUT_PATH))
    print(f"Done: {OUT_PATH}")


if __name__ == "__main__":
    convert()
